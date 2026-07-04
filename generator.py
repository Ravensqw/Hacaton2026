import json
import base64
from typing import Dict, Any

def get_file_preview(file_content: bytes, filename: str) -> dict:
    ext = filename.split('.')[-1].lower()
    
    if ext == 'csv':
        try:
            import pandas as pd
            import io
            df = pd.read_csv(io.BytesIO(file_content))
            return {
                "format": "csv",
                "columns": list(df.columns),
                "sample": df.head(3).to_dict('records'),
                "rows": len(df),
                "full_content": df.to_string()[:15000]
            }
        except Exception as e:
            return {"format": "csv", "full_content": str(file_content[:15000])}
    
    elif ext in ['xls', 'xlsx']:
        try:
            import pandas as pd
            import io
            df = pd.read_excel(io.BytesIO(file_content))
            return {
                "format": "excel",
                "columns": list(df.columns),
                "sample": df.head(3).to_dict('records'),
                "rows": len(df),
                "full_content": df.to_string()[:15000]
            }
        except Exception as e:
            return {"format": "excel", "full_content": str(file_content[:15000])}
    
    elif ext == 'pdf':
        try:
            import PyPDF2
            import io
            reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            text = ""
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            return {
                "format": "pdf",
                "pages": len(reader.pages),
                "full_content": text[:15000]
            }
        except Exception as e:
            return {"format": "pdf", "full_content": str(file_content[:15000])}
    
    elif ext == 'docx':
        try:
            import docx
            import io
            doc = docx.Document(io.BytesIO(file_content))
            text = "\n".join([para.text for para in doc.paragraphs])
            tables_text = ""
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    tables_text += row_text + "\n"
            return {
                "format": "docx",
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "full_content": (text + "\n" + tables_text)[:15000]
            }
        except Exception as e:
            return {"format": "docx", "full_content": str(file_content[:15000])}
    
    elif ext in ['png', 'jpg', 'jpeg']:
        try:
            import pytesseract
            from PIL import Image
            import io
            image = Image.open(io.BytesIO(file_content))
            text = pytesseract.image_to_string(image, lang='rus+eng')
            return {
                "format": "image",
                "size": image.size,
                "full_content": text[:15000] if text else "Текст на изображении не найден"
            }
        except Exception as e:
            return {"format": "image", "full_content": f"OCR ошибка: {str(e)[:200]}"}
    
    elif ext == 'txt':
        try:
            text = file_content.decode('utf-8', errors='ignore')
            return {
                "format": "txt",
                "full_content": text[:15000]
            }
        except:
            return {"format": "txt", "full_content": str(file_content[:15000])}
    
    elif ext == 'json':
        try:
            data = json.loads(file_content)
            return {
                "format": "json",
                "full_content": json.dumps(data, ensure_ascii=False, indent=2)[:15000]
            }
        except:
            return {"format": "json", "full_content": str(file_content[:15000])}
    
    else:
        return {
            "format": ext,
            "full_content": str(file_content[:15000]),
            "is_binary": not all(c < 128 for c in file_content[:100])
        }


class TypeScriptGenerator:
    
    def __init__(self):
        from gigachat_client import GigaChatClient
        self.gigachat = GigaChatClient()
        print("GigaChat инициализирован")
    
    async def generate(self, filename: str, file_content: bytes, json_structure: Dict) -> Dict:
        file_preview = get_file_preview(file_content, filename)
        
        ts_code = self.gigachat.generate_typescript(
            filename=filename,
            file_preview=file_preview,
            json_structure=json_structure
        )
        
        try:
            tokens_used = self.gigachat.count_tokens(ts_code)
        except:
            tokens_used = len(ts_code) // 4
        
        return {
            "ts_code": ts_code,
            "tokens_used": tokens_used,
            "preview": file_preview
        }